import zipfile
from pathlib import Path

from langchain_core.documents import Document
from pypdf import PdfReader

from app.rag.metadata_registry import normalize_source, source_metadata_for


def _load_pdf(path: Path) -> list[Document]:
    reader = PdfReader(str(path))
    docs: list[Document] = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            docs.append(Document(page_content=text, metadata={"page": i}))
    return docs


def _load_text(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8").strip()
    return [Document(page_content=text)] if text else []


def _load_docx(path: Path) -> list[Document]:
    if not zipfile.is_zipfile(path):
        return _load_text(path)
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    text = "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    return [Document(page_content=text)] if text else []


def load_file(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix in (".docx", ".doc"):
        return _load_docx(path)
    if suffix in (".md", ".markdown", ".txt"):
        return _load_text(path)
    return []


def load_directory(directory: Path) -> list[Document]:
    docs: list[Document] = []
    patterns = ("**/*.pdf", "**/*.docx", "**/*.md", "**/*.markdown", "**/*.txt")
    for pattern in patterns:
        for path in directory.glob(pattern):
            if not path.is_file():
                continue
            try:
                loaded = load_file(path)
                source = normalize_source(str(path.relative_to(directory)))
                source_meta = source_metadata_for(source)
                for doc in loaded:
                    doc.metadata["source"] = source
                    doc.metadata["file_type"] = path.suffix.lower().lstrip(".")
                    for key in ("subject", "grade", "chapter", "difficulty", "summary"):
                        value = source_meta.get(key)
                        if value:
                            doc.metadata[key] = value
                    concepts = source_meta.get("concepts") or []
                    aliases = source_meta.get("aliases") or []
                    if concepts:
                        doc.metadata["concepts"] = [str(item).strip() for item in concepts if str(item).strip()]
                    if aliases:
                        doc.metadata["aliases"] = [str(item).strip() for item in aliases if str(item).strip()]
                docs.extend(loaded)
            except Exception as exc:
                print(f"[ingest] skip {path}: {exc}")
    return docs

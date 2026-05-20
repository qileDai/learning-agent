"""生成 PDF / Word 示例知识库文件（首次运行 ingest 前执行）。"""
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

BACKEND_ROOT = Path(__file__).resolve().parent.parent
KNOWLEDGE_DIR = BACKEND_ROOT / "data" / "knowledge"


def write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    doc = DocxDocument()
    doc.add_heading(title, 0)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)


def write_pdf_simple(path: Path, title: str, lines: list[str]) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica", 14)
    c.drawString(50, y, title[:80])
    y -= 30
    c.setFont("Helvetica", 10)
    for line in lines:
        if y < 50:
            c.showPage()
            y = height - 50
            c.setFont("Helvetica", 10)
        # ASCII-safe fallback for CJK in basic PDF
        safe = line.encode("latin-1", errors="replace").decode("latin-1")
        c.drawString(50, y, safe[:90])
        y -= 14
    c.save()


def main():
    KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)

    write_docx(
        KNOWLEDGE_DIR / "biology_cell_structure.docx",
        "细胞结构与功能",
        [
            "细胞是生物体结构和功能的基本单位，分为原核细胞与真核细胞。",
            "细胞膜由磷脂双分子层构成，控制物质进出，参与信号传导。",
            "线粒体是细胞呼吸的主要场所，被称为细胞的能量工厂。",
            "叶绿体存在于植物细胞，进行光合作用，将光能转化为化学能。",
            "细胞核储存遗传物质 DNA，控制蛋白质合成与细胞分裂。",
        ],
    )

    write_docx(
        KNOWLEDGE_DIR / "english_grammar_tenses.docx",
        "英语时态学习指南",
        [
            "一般现在时：表示习惯、真理，如 Water boils at 100°C.",
            "一般过去时：表示过去发生的动作，如 I visited Beijing last year.",
            "现在进行时：表示正在进行的动作，如 She is reading a book.",
            "现在完成时：表示过去发生但与现在有联系，如 I have finished homework.",
            "被动语态：be + 过去分词，强调动作承受者，如 English is spoken worldwide.",
        ],
    )

    write_pdf_simple(
        KNOWLEDGE_DIR / "math_quadratic_equations.pdf",
        "Quadratic Equations",
        [
            "A quadratic equation: ax^2 + bx + c = 0, a != 0",
            "Discriminant D = b^2 - 4ac determines root types",
            "D > 0: two distinct real roots",
            "D = 0: one repeated real root",
            "D < 0: complex conjugate roots",
            "Formula: x = (-b +/- sqrt(D)) / (2a)",
            "Completing the square converts to vertex form",
            "Applications: projectile motion, profit optimization",
        ],
    )

    write_pdf_simple(
        KNOWLEDGE_DIR / "physics_newton_laws.pdf",
        "Newton Laws of Motion",
        [
            "First Law (Inertia): object stays at rest or uniform motion unless acted upon",
            "Second Law: F = ma, force equals mass times acceleration",
            "Third Law: every action has equal and opposite reaction",
            "Free-body diagrams help analyze forces on an object",
            "Friction: static and kinetic, f = mu * N",
            "Circular motion needs centripetal force toward center",
        ],
    )

    print(f"Generated sample docs in {KNOWLEDGE_DIR}")


if __name__ == "__main__":
    main()
